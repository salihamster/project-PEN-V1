"""
PEN Agent - Main Orchestrator

This module defines the core PENAgent class, which orchestrates the conversation
flow, memory management, and tool execution.

Refactored to use new layered memory architecture (L1, L2, L2.5, L4) + Calendar System.
"""

import os
import sys
import json
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import google.generativeai as genai
from dotenv import load_dotenv
import time

# Prompt logging directory
PROMPT_LOGS_DIR = Path(__file__).resolve().parents[2] / "prompt_logs"

# Ensure the project root is in the Python path
project_root = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(project_root))

from src.config import DATA_DIR, SERVICE_ACCOUNT_FILE, EMAIL_CONFIG, MAX_WORKERS, GEMINI_API_KEY  # type: ignore
from src.storage.data_manager import DataManager  # type: ignore
from src.agent_tools.data_tools import DataTools  # type: ignore
from src.agent_tools.email_tools import EmailTools  # type: ignore
from src.agent_tools.whatsapp_tools import WhatsAppTools  # type: ignore
from src.agent_tools.drive_tools import DriveTools  # type: ignore
from src.agent_tools.web_tools import WebTools  # type: ignore
from src.agent_tools.refresh_tools import RefreshTools  # type: ignore
from src.agent_tools.context_tools import ContextTools  # type: ignore
from src.agent_tools.calendar_tools import CalendarTools  # type: ignore
from src.agent_tools.file_system_tools import FileSystemTools # type: ignore
from src.storage.document_manager import DocumentManager # type: ignore
# L4 removed - using new layer architecture L4
from src.utils.logger import get_logger  # type: ignore
from src.utils.usage_logger import usage_logger  # type: ignore
from src.agent.tool_executor import ToolExecutor  # type: ignore
from src.agent.tool_definitions import TOOLS, FILE_SYSTEM_TOOLS  # type: ignore
from src.models.factory import ModelFactory # type: ignore
from src.models.base import BaseModelProvider # type: ignore

# Memory layers - New architecture
from layers.layer_manager import LayerManager  # type: ignore
from layers.L1 import MessageRole  # type: ignore
from layers.L4 import L4UserProfile  # type: ignore
from layers.calendar_system import CalendarSystem # type: ignore

logger = get_logger(__name__)

# ============================================================================ 
# Constants
# ============================================================================ 

DEFAULT_MODEL = "gemini-3-flash-preview"
DEFAULT_MAX_TOKENS = 24576
RATE_LIMIT_DELAY = 2  # seconds between API calls

# ============================================================================ 
# Helper Functions
# ============================================================================ 

def _to_serializable(obj):
    """
    Recursively converts Protobuf objects (RepeatedComposite, MapComposite) 
    to standard Python lists and dicts. Also handles other non-serializable types.
    """
    if isinstance(obj, (list, tuple)):
        return [_to_serializable(item) for item in obj]
    elif isinstance(obj, dict):
        return {k: _to_serializable(v) for k, v in obj.items()}
    elif hasattr(obj, "items"): # MapComposite logic
        return {k: _to_serializable(v) for k, v in obj.items()}
    elif hasattr(obj, "__iter__") and not isinstance(obj, (str, bytes)): # RepeatedComposite logic
        return [_to_serializable(item) for item in obj]
    else:
        return obj

# ============================================================================ 
# PEN Agent Class
# ============================================================================ 

class PENAgent:
    """
    Orchestrates the conversational flow, memory, and tool use for the PEN assistant.
    
    Uses new layered memory architecture:
    - L1: Active session memory (auto-persisted to JSON)
    - L2: Archived sessions
    - L2.5: Keyword-indexed search
    - L4: User profile (existing system)
    """

    def __init__(self, api_key: str, minimax_api_key: Optional[str] = None):
        if not api_key:
            raise ValueError("GEMINI_API_KEY is required.")

        genai.configure(api_key=api_key)

        # Memory Layers - New Architecture
        self.layer_manager = LayerManager()
        
        # L4 User Profile - New Architecture
        self.l4 = L4UserProfile()
        
        # Calendar System - New
        self.calendar = CalendarSystem()

        # System Prompt Template
        self.system_prompt_template = self._load_system_prompt_template()

        # Gemini Model Client (initialized on first use)
        self.client: Optional[BaseModelProvider] = None
        self.model_name = DEFAULT_MODEL

        # Tool Setup
        self.data_manager = DataManager(DATA_DIR)
        self._tool_declarations_cache = {}  # Cache for tool declarations
        self.tool_definitions = self._get_active_tool_definitions()
        self.tool_executor = self._initialize_tool_executor()
        self.valid_tool_names = [tool["name"] for tool in TOOLS]

        # Conversation History for the current session (for the model)
        # Load existing L1 messages into Gemini conversation history
        self.messages: List[Any] = []
        self._load_l1_into_conversation()

    def _get_active_tool_definitions(self) -> List[Dict[str, Any]]:
        """Returns the list of currently active tool definitions as raw dicts."""
        fs_active = self.layer_manager.l1.session_metadata.system_state.get("file_system_tools_active", False)
        
        active_tools = TOOLS[:]
        if fs_active:
            active_tools.extend(FILE_SYSTEM_TOOLS)
        return active_tools

    def _load_l1_into_conversation(self) -> None:
        """Load existing L1 messages into Gemini conversation history."""
        try:
            l1_messages = self.layer_manager.l1.get_all_messages()
            
            for msg in l1_messages:
                try:
                    role = msg.get("role")
                    content = msg.get("content", "")
                    
                    # Validate content is a string
                    if not isinstance(content, str):
                        logger.warning(f"Skipping message with non-string content: {type(content)}")
                        continue
                    
                    # Skip empty messages
                    if not content or not content.strip():
                        continue
                    
                    # Convert L1 roles to Gemini roles
                    if role == "user":
                        self.messages.append({"role": "user", "parts": [content]})
                    elif role == "assistant":
                        self.messages.append({"role": "model", "parts": [content]})
                    elif role == "system":
                        # Gemini doesn't have a dedicated "system" role in history (only in system_instruction)
                        # We inject system events as "user" messages with a special tag to simulate system prompts
                        self.messages.append({"role": "user", "parts": [f"{content}"]})
                    # Skip tool messages for now (they need special handling)
                except Exception as msg_error:
                    logger.warning(f"Failed to load individual message: {msg_error}")
                    continue
            
            if l1_messages:
                logger.info(f"Loaded {len(self.messages)} valid messages from {len(l1_messages)} L1 messages")
        except Exception as e:
            logger.warning(f"Could not load L1 messages: {e}")

    def _load_system_prompt_template(self) -> str:
        """Loads the system prompt from a file."""
        return """You are PENNY - the user's work colleague. Natural, concise, effective communication.

## CORE RULES
1. Don't show thinking process, give direct results
2. Short question = short answer, complex task = use tools + summarize
3. Use tools silently (don't say \"Looking...\" or \"Searching...\")
4. Create/update context when learning new info (no permission needed)
5. Ask before deleting or major changes
6. ALWAYS respond to user after using tools
7. **PERSISTENT MEMORY:** If user says "remember this", "from now on", or sets a preference, YOU MUST use `add_behavioral_directive` or `create_context`. Do not just say "I will remember".

## PEN WORKSPACE & FILE SYSTEM
You have access to a secure file system (PEN WorkSpace) in `data/user_docs`.
- **Artifacts:** When you create or edit a file (using `write_to_file`), it appears in the user's "Artifact Panel". Use this for:
  - Drafting emails/reports
  - Coding
  - Detailed plans
  - Meeting notes
- **Workflow:**
  - "Daily Plan": Check/Create `daily_plans/YYYY-MM-DD.md`
  - "Meeting Notes": Create `notes/meeting_[topic]_[date].md` and link to calendar event.
  - "Project Docs": Store in `projects/[project_name]/`.
- **Linking:** Always try to link created documents to relevant Calendar Events using `create_event(..., linked_file=path)`.

## TOOL STRATEGY
- **CRITICAL DISTINCTION:**
  - If user asks "What did **WE** talk about?", "What did **I** tell **YOU**?", "Recall our last session": -> **INTERNAL MEMORY**. Use `search_memory` then `read_archived_session`. DO NOT check WhatsApp.
  - If user asks "What did I tell **Ali**?", "Check my messages": -> **EXTERNAL SOURCES**. Use `search_messages` / `get_whatsapp_messages`.
- **Memory/Rules:** Use `add_behavioral_directive` for rules ("don't use emojis").
- Search info: search_messages → get_whatsapp_messages → search_web
- Find context: search_contexts → read_context
- New info: create_context or update_context
- Calendar/Plans: calendar_tools_open -> read_calendar/create_event (link files if needed)
- Files: read_file, write_to_file, list_files (User Docs)
- Check multiple sources, present the best
- Summarize tool results and inform user

## TONE
- Work colleague, not corporate assistant
- Natural language, technical when needed
- Dry humor OK, no emojis

## ACTIVE CONTEXT
**Date:** {current_time}

{l4_context}

{episodic_context}

{tool_context}
"""

    def _initialize_tool_executor(self) -> ToolExecutor:
        """Initializes all tool classes and the ToolExecutor."""
        data_tools = DataTools(self.data_manager)
        email_tools = EmailTools(self.data_manager)
        whatsapp_tools = WhatsAppTools(self.data_manager)
        brave_api_key = os.getenv("BRAVE_API_KEY")
        web_tools = WebTools(brave_api_key=brave_api_key)
        refresh_tools = RefreshTools(
            data_manager=self.data_manager,
            email_config=EMAIL_CONFIG,
            service_account_file=str(SERVICE_ACCOUNT_FILE) if SERVICE_ACCOUNT_FILE.exists() else None
        )
        context_tools = ContextTools(self.l4)
        calendar_tools = CalendarTools(self.calendar)
        
        # Initialize File System Tools
        doc_manager = DocumentManager(DATA_DIR / "user_docs")
        file_system_tools = FileSystemTools(doc_manager)
        
        drive_tools: Optional[DriveTools] = None
        if SERVICE_ACCOUNT_FILE.exists():
            try:
                drive_tools = DriveTools(str(SERVICE_ACCOUNT_FILE), folder_name="Wpmesages")
            except Exception as e:
                logger.error(f"Failed to initialize DriveTools: {e}")

        return ToolExecutor(
            data_tools=data_tools,
            email_tools=email_tools,
            whatsapp_tools=whatsapp_tools,
            web_tools=web_tools,
            refresh_tools=refresh_tools,
            context_tools=context_tools,
            calendar_tools=calendar_tools,
            file_system_tools=file_system_tools,
            drive_tools=drive_tools,
            l1_layer=self.layer_manager.l1,  # For context management tool
            layer_manager=self.layer_manager, # For L2/L2.5 access
        )

    def _get_system_prompt(self) -> str:
        """Constructs the final system prompt with dynamic context."""
        from datetime import datetime
        
        # L4 Context - Use RAW profile data (No narrative summary)
        l4_data = self.l4.load_profile()
        user_profile = l4_data.get("user_profile", {})
        
        # Build L4 Context String
        l4_context_parts = ["=== USER PROFILE ==="]
        
        # 1. Identity
        name = user_profile.get("name", "User")
        l4_context_parts.append(f"Name: {name}")
        
        # 2. Behavioral Directives (CRITICAL - RAW)
        directives = user_profile.get("behavioral_directives", [])
        if directives:
            l4_context_parts.append("\n=== OPERATING INSTRUCTIONS (Behavioral Directives) ===")
            for d in directives:
                l4_context_parts.append(f"• {d}")
                
        # 3. Profile Facts (Lists)
        categories = {
            "Goals": "goals",
            "Current Projects": "projects",
            "Preferences": "preferences",
            "Expertise": "expertise",
            "Interests": "interests",
            "Notes": "notes"
        }
        
        l4_context_parts.append("\n=== USER FACTS ===")
        has_facts = False
        for label, key in categories.items():
            items = user_profile.get(key, [])
            if items:
                has_facts = True
                l4_context_parts.append(f"\n{label}:")
                for item in items:
                    l4_context_parts.append(f"- {item}")
        
        if not has_facts:
            l4_context_parts.append("(No facts recorded yet)")
        
        # Get Available Context Headers (Knowledge Base)
        headers = self.l4.get_contexts_headers()
        l4_context_parts.append("\n=== AVAILABLE KNOWLEDGE BASE (Contexts) ===")
        if headers:
            for h in headers[:15]: # Limit to top 15 recent contexts
                l4_context_parts.append(f"• [{h['id']}] {h['title']} ({h['type']})")
            if len(headers) > 15:
                l4_context_parts.append(f"... and {len(headers)-15} more (use context_tools_open or search_contexts)")
        else:
            l4_context_parts.append("(No active knowledge contexts)")
            
        l4_context = "\n".join(l4_context_parts)

        # L2.5 Context (Recent Episodes)
        recent_sessions = self.layer_manager.l2_5.get_recent_sessions(limit=3)
        episodic_context = ""
        if recent_sessions:
            episodic_context = "=== RECENT EPISODES (L2.5) ===\n"
            for session_data in recent_sessions:
                date = session_data.get("created_at", "")[:10]  # YYYY-MM-DD
                summary = session_data.get("summary", "")
                episodic_context += f"- {date}: {summary}\n"

        current_time = datetime.now().strftime("%d %B %Y, %H:%M")
        
        # Tool Context (L1 tool outputs - expanded/collapsed)
        tool_context = self.layer_manager.l1.get_tool_context_for_model()
        
        # Dynamic Workspace Context
        workspace_context = ""
        fs_active = self.layer_manager.l1.session_metadata.system_state.get("file_system_tools_active", False)
        if fs_active:
            workspace_context = """
## PEN WORKSPACE & FILE SYSTEM (ACTIVE)
You have access to a secure file system (PEN WorkSpace) in `data/user_docs`.
- **Artifacts:** When you create or edit a file (using `write_to_file`), it appears in the user's "Artifact Panel".
- **Workflow:**
  - "Daily Plan": Check/Create `daily_plans/YYYY-MM-DD.md`
  - "Meeting Notes": Create `notes/meeting_[topic]_[date].md` and link to calendar event.
  - "Project Docs": Store in `projects/[project_name]/`.
- **Linking:** Always try to link created documents to relevant Calendar Events using `create_event(..., linked_file=path)`.
"""

        full_prompt = self.system_prompt_template.format(
            current_time=current_time,
            l4_context=l4_context,
            episodic_context=episodic_context.strip(),
            tool_context=tool_context
        )
        
        if workspace_context:
            # Inject workspace context before TOOL STRATEGY
            if "## TOOL STRATEGY" in full_prompt:
                full_prompt = full_prompt.replace("## TOOL STRATEGY", workspace_context + "\n## TOOL STRATEGY")
            else:
                full_prompt += workspace_context
                
        return full_prompt

    def _clean_conversation_history(self):
        """
        Cleans self.messages to remove tool outputs/calls from previous turns.
        We rely on the System Prompt (managed by L1) to provide tool context.
        This prevents unbounded context growth in the active session list.
        """
        cleaned_messages = []
        for msg in self.messages:
            # Handle Dict messages
            if isinstance(msg, dict):
                parts = msg.get("parts", [])
                should_skip = False
                
                if isinstance(parts, list):
                    for part in parts:
                        # Check for function response (Tool Output)
                        if hasattr(part, "function_response") and part.function_response:
                            should_skip = True
                            break
                        if isinstance(part, dict) and ("functionResponse" in part or "function_response" in part):
                             should_skip = True
                             break
                
                if not should_skip:
                    cleaned_messages.append(msg)
                    
            # Handle Protobuf Content objects (Model Responses)
            elif hasattr(msg, "parts"):
                new_parts = []
                for part in msg.parts:
                    # Skip function calls (Model side)
                    if hasattr(part, "function_call") and part.function_call:
                        continue 
                    # Skip function responses (User side - if wrapped in Content)
                    if hasattr(part, "function_response") and part.function_response:
                        continue 
                    new_parts.append(part)
                
                # Only keep message if it has remaining parts (e.g. text)
                if new_parts:
                    if len(new_parts) == len(msg.parts):
                        cleaned_messages.append(msg)
                    else:
                        # Reconstruct as dict to avoid protobuf issues
                        cleaned_messages.append({"role": msg.role, "parts": new_parts})
            
            else:
                cleaned_messages.append(msg)
        
        if len(self.messages) != len(cleaned_messages):
            logger.info(f"Context cleanup: Reduced history from {len(self.messages)} to {len(cleaned_messages)} messages (removed tool outputs)")
            self.messages = cleaned_messages

    def set_model(self, model_name: str):
        """Change the model used by the agent."""
        if model_name != self.model_name:
            logger.info(f"Changing model from {self.model_name} to {model_name}")
            self.model_name = model_name
            self.client = None
            self._ensure_client()

    def _ensure_client(self):
        """Ensures the model provider is initialized with the latest system prompt."""
        final_prompt = self._get_system_prompt()
        
        # Log prompt length
        prompt_length = len(final_prompt)
        logger.debug(f"System prompt length: {prompt_length} chars (~{prompt_length // 4} tokens)")

        if self.client is None or self.client.config.model_id != self.model_name:
            logger.info(f"==== Initializing model provider: {self.model_name} ====")
            try:
                self.client = ModelFactory.create_provider(self.model_name)
            except Exception as e:
                logger.error(f"Failed to create provider: {e}")
                # Fallback if creation fails
                if self.model_name != DEFAULT_MODEL:
                    self.model_name = DEFAULT_MODEL
                    self.client = ModelFactory.create_provider(DEFAULT_MODEL)
                else:
                    raise e
        
        # Always re-initialize with latest prompt and tools
        self.tool_definitions = self._get_active_tool_definitions()
        self.client.initialize(final_prompt, self.tool_definitions)

    def chat(self, user_message: str, files: list = None) -> str:
        """Main chat loop for processing user input."""
        
        # 1. Run TTL Logic (on previous history)
        collapsed_ids = self.layer_manager.l1.tick_ttl()
        if collapsed_ids:
            logger.info(f"Auto-collapsed {len(collapsed_ids)} tool outputs due to TTL expiry")

        # 2. Clean Message History (remove old tool outputs)
        self._clean_conversation_history()

        # 3. Ensure Client (Load System Prompt with updated L1 state)
        self._ensure_client()
        assert self.client is not None, "Client should be initialized"

        # ✅ SMART REMINDERS: Check for upcoming events on first message
        is_first_message = len(self.messages) == 0
        if is_first_message:
            smart_reminder = self._get_smart_reminders()
            if smart_reminder:
                # Inject as system message (model will see it)
                self.messages.append({"role": "user", "parts": [smart_reminder]})
                logger.info(f"Added smart reminder: {smart_reminder[:100]}...")

        # Build message parts
        parts = [user_message]
        file_metadata = []
        
        if files:
            for f in files:
                file_name = f.get("name", "unknown")
                file_type = f.get("type", "")
                file_content = f.get("content", "")
                
                if file_type.startswith("image/"):
                    import base64
                    if isinstance(file_content, bytes):
                        b64_data = base64.b64encode(file_content).decode("utf-8")
                    else:
                        b64_data = file_content
                    parts.append(genai.protos.Part(
                        inline_data={"mime_type": file_type, "data": b64_data}
                    ))
                    file_metadata.append({"name": file_name, "type": file_type, "size": len(file_content), "content": b64_data, "is_image": True})
                else:
                    text_content = ""
                    # Handle specific file types
                    if file_type == "application/pdf" or file_name.lower().endswith(".pdf"):
                         text_content = self._extract_pdf_text(file_content)
                    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.lower().endswith(".docx"):
                         text_content = self._extract_docx_text(file_content)
                    else:
                        # Fallback for text files
                        if isinstance(file_content, bytes):
                            text_content = file_content.decode("utf-8", errors="replace")
                        else:
                            text_content = file_content
                    
                    parts.append(f"[ATTACHED FILE: {file_name}]\n{text_content}\n[END FILE: {file_name}]")
                    preview_content = text_content[:50000] if len(text_content) > 50000 else text_content
                    file_metadata.append({"name": file_name, "type": file_type, "size": len(text_content), "content": preview_content})

        self.layer_manager.add_user_message(user_message, metadata={"files": file_metadata} if file_metadata else None)
        self.messages.append({"role": "user", "parts": parts})
        
        logger.info(f"=== USER MESSAGE RECEIVED ({len(user_message)} chars) ===")
        if len(user_message) > 1000:
            logger.info(user_message[:1000] + "...")
        else:
            logger.info(user_message)
        
        retry_count = 0
        max_retries = 3
        request_start_time = time.time()
        all_tool_calls: List[str] = []
        
        while True:
            try:
                time.sleep(RATE_LIMIT_DELAY)
                self._log_prompt_to_file()

                try:
                    # Use abstract generate method
                    response_dict = self.client.generate(self.messages)
                except Exception as api_error:
                    logger.error(f"API Error: {api_error}")
                    if retry_count < max_retries:
                        retry_count += 1
                        logger.info(f"Retrying... (attempt {retry_count})")
                        self.client = None
                        self._ensure_client()
                        continue
                    else:
                        raise api_error

                text_response = response_dict.get("content", "")
                tool_calls = response_dict.get("tool_calls", [])
                raw_response = response_dict.get("raw_response")

                # If no tools, process text response
                if not tool_calls:
                    if text_response.strip():
                        if hasattr(self, '_last_response_hash'):
                            current_hash = hash(text_response.strip()[:500])
                            if current_hash == self._last_response_hash:
                                logger.warning("Duplicate response detected, skipping")
                                return text_response
                        self._last_response_hash = hash(text_response.strip()[:500])
                        
                        print()
                        self.layer_manager.add_assistant_message(text_response)
                        self.messages.append({"role": "model", "parts": [text_response]})
                        
                        response_time_ms = int((time.time() - request_start_time) * 1000)
                        session_source = getattr(self, 'session_source', 'cli')
                        usage_logger.log_request(
                            model=self.model_name,
                            provider=self.client.config.provider,
                            prompt_tokens=0, 
                            completion_tokens=0,
                            response_time_ms=response_time_ms,
                            tool_calls=all_tool_calls,
                            session_source=session_source,
                            success=True
                        )
                        return text_response
                    else:
                        return "I'm sorry, I couldn't generate a response."

                # Add model's response (with tool calls) to history
                if self.client.config.provider == "gemini":
                    self.messages.append(raw_response)
                else:
                    self.messages.append({"role": "assistant", "content": text_response, "tool_calls": tool_calls})
                
                tool_results = []
                tool_names = []
                
                for tool_call in tool_calls:
                    tool_name = tool_call.get("name")
                    tool_input = tool_call.get("arguments", {})
                    tool_id = tool_call.get("id") # For Anthropic
                    
                    tool_names.append(tool_name)
                    
                    try:
                        # Handle State Changes
                        if tool_name == "file_system_tools_open":
                            self.layer_manager.l1.session_metadata.system_state["file_system_tools_active"] = True
                            self.layer_manager.l1.save_to_file()
                            self._ensure_client() 

                        self._emit_tool_event(tool_name, "running")
                        result_str = self.tool_executor.execute(tool_name, tool_input)
                        self._emit_tool_event(tool_name, "completed")
                        
                        try:
                            result_dict = json.loads(result_str)
                        except:
                            result_dict = {"result": result_str}
                        
                        self.layer_manager.add_tool_call(
                            tool_name=tool_name,
                            tool_input=tool_input,
                            tool_output=result_dict
                        )
                        
                        if self.client.config.provider == "anthropic":
                             result_dict["tool_use_id"] = tool_id
                             
                        formatted_result = self.client.format_tool_result(tool_name, result_dict)
                        tool_results.append(formatted_result)

                    except Exception as tool_err:
                        logger.error(f"Tool error {tool_name}: {tool_err}")
                        error_res = {"error": str(tool_err)}
                        if self.client.config.provider == "anthropic":
                             error_res["tool_use_id"] = tool_id
                        tool_results.append(self.client.format_tool_result(tool_name, error_res))

                if tool_names:
                    tools_str = ", ".join(tool_names)
                    print(f"🔧 [{tools_str}]", end=" ", flush=True)
                    
                    if self.client.config.provider == "gemini":
                        self.messages.append({"role": "user", "parts": tool_results})
                    else:
                        for res in tool_results:
                            self.messages.append({"role": "user", "content": [res] if isinstance(res, dict) else res})
                    
                    all_tool_calls.extend(tool_names)
                else:
                    return "I'm sorry, I couldn't generate a response."
            
            except Exception as e:
                import traceback
                error_details = traceback.format_exc()
                logger.error(f"Error in chat loop: {e}\n{error_details}")
                return f"An error occurred: {str(e)}"

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text.strip() or "PDF'den metin cikarildi (bos)"
        except ImportError:
            return "PDF okuma icin PyMuPDF (fitz) kurulu degil"
        except Exception as e:
            return f"PDF okunamadi: {str(e)}"

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            import docx
            import io
            
            # Load the docx from bytes
            doc = docx.Document(io.BytesIO(content))
            
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    if any(row_text):
                        full_text.append(" | ".join(row_text))
                    
            return "\n".join(full_text)
        except ImportError:
            return "DOCX okuma icin python-docx kurulu degil"
        except Exception as e:
            return f"DOCX okunamadi: {str(e)}"


    def sleep(self):
        """Archives the current session."""
        logger.info("Initiating sleep cycle...")
        try:
            l4_result = self.update_l4_from_conversation()
            if l4_result.get("status") == "success":
                logger.info(f"L4 updated: {l4_result.get('updates_count', 0)} new insights")
        except Exception as e:
            logger.error(f"L4 update error: {e}")
        
        try:
            # New Consolidation Logic
            consolidation_result = self.l4.consolidate_profile_data()
            logger.info(f"Profile consolidation: {consolidation_result.get('message', 'done')}")
        except Exception as e:
            logger.error(f"Profile consolidation error: {e}")
        
        result = self.layer_manager.trigger_sleep_cycle()
        if result.get("status") == "success":
            logger.info(f"Sleep cycle completed: {result.get('message')}")
        else:
            logger.error(f"Sleep cycle failed: {result.get('message')}")
        
        self.reset(clear_l1=False)

    def reset(self, clear_l1: bool = True):
        """Resets conversation history."""
        self.messages = []
        if clear_l1:
            self.layer_manager.l1.clear_session()
            logger.info("Agent conversation history and L1 session cleared.")
        else:
            logger.info("Agent conversation history reset (L1 preserved).")

    def _log_prompt_to_file(self):
        """Log prompt to file."""
        try:
            PROMPT_LOGS_DIR.mkdir(exist_ok=True)
            timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            filename = f"prompt_{timestamp}.json"
            filepath = PROMPT_LOGS_DIR / filename
            system_instruction = self._get_system_prompt()
            
            # Simplified logging logic for brevity
            prompt_data = {
                "timestamp": datetime.now().isoformat(),
                "model": self.model_name,
                "system_instruction": system_instruction,
                "messages_count": len(self.messages)
            }
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(prompt_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"Failed to log prompt: {e}")

    def _emit_tool_event(self, tool_name: str, status: str):
        """Emit tool event."""
        try:
            from web_server import emit_tool_event_sync
            emit_tool_event_sync(tool_name, status)
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"Could not emit tool event: {e}")

    def update_l4_from_conversation(self) -> Dict[str, Any]:
        """Updates L4 user profile."""
        try:
            l1_session_context = self.layer_manager.l1.get_session_context()
            if not l1_session_context.get("messages"):
                return {"status": "no_messages"}
            return self.l4.update_profile_from_session(l1_session_context)
        except Exception as e:
            logger.error(f"L4 update error: {e}", exc_info=True)
            return {"status": "error", "message": str(e)}
    
    def _get_smart_reminders(self) -> str:
        """
        Generate smart reminders for upcoming events.
        Called on first message of a session.
        
        Returns:
            Formatted reminder string or empty string if no reminders
        """
        from datetime import datetime, timedelta
        
        try:
            now = datetime.now()
            reminders = []
            
            # Load calendar data
            data = self.calendar._load_data()
            events = data.get("events", {})
            
            if not events:
                return ""
            
            for event_data in events.values():
                # Skip inactive events
                if event_data.get("status") != "active":
                    continue
                
                # Get start time
                start_time_str = event_data.get("start_time")
                if not start_time_str:
                    continue
                
                try:
                    start_time = datetime.fromisoformat(start_time_str)
                    time_diff = start_time - now
                    
                    # Skip past events
                    if time_diff.total_seconds() < 0:
                        continue
                    
                    title = event_data.get("title", "Event")
                    
                    # TODAY (0-24 hours)
                    if timedelta(0) <= time_diff <= timedelta(hours=24):
                        hours_left = int(time_diff.total_seconds() / 3600)
                        minutes_left = int((time_diff.total_seconds() % 3600) / 60)
                        
                        # Format time
                        if hours_left > 0:
                            time_str = f"{hours_left}h"
                            if minutes_left > 0:
                                time_str += f" {minutes_left}m"
                        else:
                            time_str = f"{minutes_left}m"
                        
                        # Categorize by type
                        title_lower = title.lower()
                        if "sınav" in title_lower or "exam" in title_lower:
                            reminders.append(f"⚠️ EXAM TODAY: {title} in {time_str}")
                        elif "toplantı" in title_lower or "meeting" in title_lower:
                            reminders.append(f"📞 MEETING TODAY: {title} in {time_str}")
                        elif "çalış" in title_lower or "study" in title_lower:
                            reminders.append(f"📚 STUDY SESSION: {title} in {time_str}")
                        else:
                            reminders.append(f"📅 TODAY: {title} in {time_str}")
                    
                    # TOMORROW (24-48 hours)
                    elif timedelta(hours=24) <= time_diff <= timedelta(hours=48):
                        title_lower = title.lower()
                        if "sınav" in title_lower or "exam" in title_lower:
                            reminders.append(f"📚 EXAM TOMORROW: {title} - Prepare now!")
                        elif "toplantı" in title_lower or "meeting" in title_lower:
                            reminders.append(f"📞 MEETING TOMORROW: {title}")
                    
                    # THIS WEEK (2-7 days)
                    elif timedelta(days=2) <= time_diff <= timedelta(days=7):
                        days_left = int(time_diff.total_seconds() / 86400)
                        title_lower = title.lower()
                        if "sınav" in title_lower or "exam" in title_lower:
                            reminders.append(f"📖 EXAM IN {days_left} DAYS: {title}")
                
                except (ValueError, TypeError):
                    continue
            
            # Format reminders
            if reminders:
                # Limit to top 3 most urgent
                reminders = reminders[:3]
                reminder_text = "[SYSTEM REMINDER] " + " | ".join(reminders)
                return reminder_text
            
            return ""
        
        except Exception as e:
            logger.error(f"Error generating smart reminders: {e}")
            return ""