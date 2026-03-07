"""Minimal(!) web server for PENAgent.

Runs a simple HTTP API to expose pen_agent.PENAgent as a chat endpoint:
- POST /api/pen/chat {"message": "..."}
- POST /api/pen/reset
- GET /api/pen/events (SSE for real-time tool activity)

And serves the static web UI from ./web (index.html, style.css, script.js).
"""

from __future__ import annotations

import os
import json
import asyncio
import threading
import queue as thread_queue
from pathlib import Path
from typing import List, Dict, Any, Optional
from collections import deque
from datetime import datetime

from fastapi import FastAPI, HTTPException, Query, File, UploadFile, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.routing import Mount
from pydantic import BaseModel
import tempfile
import base64

from src.agent.agent import PENAgent
from src.utils.logger import get_logger
from src.storage.document_manager import DocumentManager
from src.config import DATA_DIR

logger = get_logger(__name__)

# Initialize Document Manager
doc_manager = DocumentManager(DATA_DIR / "user_docs")

# Global event queue for SSE - Thread-safe implementation
class EventBroadcaster:
    def __init__(self):
        self.subscribers: List[asyncio.Queue] = []
        self._pending_events: thread_queue.Queue = thread_queue.Queue()
        self._lock = threading.Lock()
    
    async def subscribe(self) -> asyncio.Queue:
        async_queue: asyncio.Queue = asyncio.Queue()
        with self._lock:
            self.subscribers.append(async_queue)
        logger.info(f"SSE subscriber added. Total: {len(self.subscribers)}")
        return async_queue
    
    def unsubscribe(self, queue: asyncio.Queue):
        with self._lock:
            if queue in self.subscribers:
                self.subscribers.remove(queue)
        logger.info(f"SSE subscriber removed. Total: {len(self.subscribers)}")
    
    async def broadcast(self, event_type: str, data: Dict[str, Any]):
        event = {"type": event_type, "data": data, "timestamp": datetime.now().isoformat()}
        logger.debug(f"Broadcasting event: {event_type} to {len(self.subscribers)} subscribers")
        with self._lock:
            subscribers_copy = self.subscribers[:]
        for async_queue in subscribers_copy:
            try:
                async_queue.put_nowait(event)
            except asyncio.QueueFull:
                logger.warning("SSE queue full, dropping event")
            except Exception as e:
                logger.warning(f"Failed to broadcast to subscriber: {e}")
    
    def broadcast_sync(self, event_type: str, data: Dict[str, Any]):
        """Synchronous broadcast - puts events directly into async queues from sync context"""
        event = {"type": event_type, "data": data, "timestamp": datetime.now().isoformat()}
        logger.debug(f"Sync broadcasting event: {event_type} - {data.get('tool_name', 'unknown')} to {len(self.subscribers)} subscribers")
        
        with self._lock:
            subscribers_copy = self.subscribers[:]
        
        for async_queue in subscribers_copy:
            try:
                # put_nowait is thread-safe for asyncio.Queue
                async_queue.put_nowait(event)
            except Exception as e:
                logger.warning(f"Failed to sync broadcast: {e}")

event_broadcaster = EventBroadcaster()

# Function to emit tool events (will be called from agent)
async def emit_tool_event(tool_name: str, status: str, details: Dict[str, Any] = None):
    await event_broadcaster.broadcast("tool_activity", {
        "tool_name": tool_name,
        "status": status,
        "details": details or {}
    })

def emit_tool_event_sync(tool_name: str, status: str, details: Dict[str, Any] = None):
    """Synchronous wrapper for emitting tool events - uses direct queue access"""
    logger.info(f"emit_tool_event_sync called: {tool_name} - {status}")
    event_broadcaster.broadcast_sync("tool_activity", {
        "tool_name": tool_name,
        "status": status,
        "details": details or {}
    })

BASE_DIR = Path(__file__).parent
WEB_DIR = BASE_DIR / "web"

app = FastAPI(title="PEN Agent Web API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class ChatRequest(BaseModel):
    message: str


class ToolRecord(BaseModel):
    name: str
    status: str
    parameters: Dict[str, Any] | None = None
    result: Any | None = None
    duration_ms: int | None = None
    started_at: str | None = None


# Single PENAgent instance for the process
_agent_instance: PENAgent | None = None


def get_agent() -> PENAgent:
    global _agent_instance
    if _agent_instance is None:
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise RuntimeError("GEMINI_API_KEY is not set in environment")
        minimax = os.getenv("MINIMAX_API_KEY")
        _agent_instance = PENAgent(api_key=api_key, minimax_api_key=minimax)
        _agent_instance.session_source = "web"  # Set session source for web
    return _agent_instance


async def process_attached_files(files_data: list, agent) -> str:
    """Process uploaded files and return context string"""
    results = []
    
    logger.info(f"=== PROCESSING {len(files_data)} FILES ===")
    
    for idx, file_info in enumerate(files_data):
        name = file_info["name"]
        mime_type = file_info["type"]
        content = file_info["content"]
        
        logger.info(f"[File {idx+1}] Name: {name}, MIME: {mime_type}, Size: {len(content)} bytes")
        
        try:
            if mime_type.startswith("image/"):
                # Use Gemini Vision for images
                result = await process_image_file(content, mime_type, name)
                results.append(f"[Gorsel: {name}]\n{result}")
            elif mime_type == "application/pdf":
                # Extract text from PDF
                result = extract_pdf_text(content)
                results.append(f"[PDF: {name}]\n{result[:5000]}...")
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or name.endswith(".docx"):
                # Extract text from DOCX
                result = extract_docx_text(content)
                results.append(f"[DOCX: {name}]\n{result[:5000]}...")
            elif mime_type.startswith("text/") or name.endswith((".txt", ".md", ".json", ".csv", ".py", ".js", ".html", ".css")):
                # Text files - decode directly
                text = content.decode("utf-8", errors="replace")
                logger.info(f"[File {idx+1}] Text decoded OK, length: {len(text)} chars")
                logger.info(f"[File {idx+1}] First 200 chars: {text[:200]}")
                results.append(f"--- DOSYA: {name} ---\n{text[:5000]}\n--- DOSYA SONU: {name} ---")
            else:
                # Try to decode as text anyway
                try:
                    text = content.decode("utf-8", errors="replace")
                    logger.info(f"[File {idx+1}] Unknown type decoded as text, length: {len(text)} chars")
                    results.append(f"--- DOSYA: {name} ---\n{text[:5000]}\n--- DOSYA SONU: {name} ---")
                except:
                    results.append(f"[Dosya: {name}] (Desteklenmeyen format: {mime_type})")
        except Exception as e:
            logger.error(f"Error processing file {name}: {e}")
            results.append(f"[Dosya: {name}] (Islenemedi: {str(e)})")
    
    final_result = "\n\n".join(results)
    logger.info(f"=== FILE PROCESSING COMPLETE ===")
    logger.info(f"Total results: {len(results)} files processed")
    logger.info(f"Combined output length: {len(final_result)} chars")
    return final_result


async def process_image_file(content: bytes, mime_type: str, filename: str) -> str:
    """Process image using Gemini Vision API"""
    try:
        import google.generativeai as genai
        
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            return "Gorsel isleme icin Gemini API anahtari gerekli"
        
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.0-flash")
        
        # Create image part
        image_part = {
            "mime_type": mime_type,
            "data": base64.b64encode(content).decode("utf-8")
        }
        
        response = model.generate_content([
            "Bu gorseli detayli bir sekilde analiz et ve icerigini acikla. Turkce yaz.",
            image_part
        ])
        
        return response.text
    except Exception as e:
        logger.error(f"Vision API error: {e}")
        return f"Gorsel analiz edilemedi: {str(e)}"


def extract_pdf_text(content: bytes) -> str:
    """Extract text from PDF using PyMuPDF if available"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip() or "PDF'den metin cikarildi (bos)"
    except ImportError:
        return "PDF okuma icin PyMuPDF (fitz) kurulu degil"
    except Exception as e:
        return f"PDF okunamadi: {str(e)}"


def extract_docx_text(content: bytes) -> str:
    """Extract text from DOCX using python-docx"""
    try:
        import docx
        import io
        
        # Load the docx from bytes
        doc = docx.Document(io.BytesIO(content))
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                if any(row_text):
                    full_text.append(" | ".join(row_text))
                
        return "\n".join(full_text)
    except ImportError:
        return "DOCX okuma icin python-docx kurulu degil"
    except Exception as e:
        return f"DOCX okunamadi: {str(e)}"


@app.get("/api/pen/calendar/events")
async def get_calendar_events():
    """Get all calendar events grouped by date"""
    try:
        agent = get_agent()
        calendar_system = agent.calendar
        
        # Load calendar data
        data = calendar_system._load_data()
        events = data.get("events", {})
        
        # Group events by date
        events_by_date = {}
        for event_id, event_data in events.items():
            # Get event type
            event_type = event_data.get("type", "fixed")
            
            # Determine date based on event type
            if event_type == "windowed":
                # For windowed events, use window_start
                date_str = event_data.get("window_start", "")
            else:
                # For fixed events, use start_time
                date_str = event_data.get("start_time", "")
            
            # Skip events without dates
            if not date_str:
                continue
            
            # Extract date key (YYYY-MM-DD)
            date_key = date_str.split("T")[0]
            
            # For windowed events, add to all dates in the window
            if event_type == "windowed" and event_data.get("window_end"):
                window_end = event_data.get("window_end", "")
                if window_end:
                    # Parse dates
                    from datetime import datetime, timedelta
                    try:
                        start_dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
                        end_dt = datetime.fromisoformat(window_end.replace("Z", "+00:00"))
                        
                        # Add event to each day in the window
                        current_dt = start_dt
                        while current_dt <= end_dt:
                            day_key = current_dt.strftime("%Y-%m-%d")
                            if day_key not in events_by_date:
                                events_by_date[day_key] = []
                            
                            # Only add once per day
                            if not any(e["id"] == event_id for e in events_by_date[day_key]):
                                events_by_date[day_key].append({
                                    "id": event_id,
                                    "title": event_data.get("title", ""),
                                    "type": event_type,
                                    "start_date": event_data.get("start_time"),
                                    "end_date": event_data.get("end_time"),
                                    "window_start": event_data.get("window_start"),
                                    "window_end": event_data.get("window_end"),
                                    "description": event_data.get("description", ""),
                                    "tags": event_data.get("tags", []),
                                    "status": event_data.get("status", "active"),
                                    "duration_minutes": event_data.get("duration_minutes")
                                })
                            
                            current_dt += timedelta(days=1)
                    except Exception as e:
                        logger.error(f"Error parsing windowed event dates: {e}")
                        # Fallback: just add to start date
                        if date_key not in events_by_date:
                            events_by_date[date_key] = []
                        events_by_date[date_key].append({
                            "id": event_id,
                            "title": event_data.get("title", ""),
                            "type": event_type,
                            "start_date": event_data.get("start_time"),
                            "end_date": event_data.get("end_time"),
                            "window_start": event_data.get("window_start"),
                            "window_end": event_data.get("window_end"),
                            "description": event_data.get("description", ""),
                            "tags": event_data.get("tags", []),
                            "status": event_data.get("status", "active"),
                            "duration_minutes": event_data.get("duration_minutes")
                        })
            else:
                # Fixed events: add to single date
                if date_key not in events_by_date:
                    events_by_date[date_key] = []
                
                events_by_date[date_key].append({
                    "id": event_id,
                    "title": event_data.get("title", ""),
                    "type": event_type,
                    "start_date": event_data.get("start_time"),
                    "end_date": event_data.get("end_time"),
                    "window_start": event_data.get("window_start"),
                    "window_end": event_data.get("window_end"),
                    "description": event_data.get("description", ""),
                    "tags": event_data.get("tags", []),
                    "status": event_data.get("status", "active"),
                    "duration_minutes": event_data.get("duration_minutes")
                })
        
        return {
            "status": "success",
            "events": events_by_date
        }
    except Exception as e:
        logger.error(f"Error getting calendar events: {e}", exc_info=True)
        return {"status": "error", "error": str(e), "events": {}}


@app.post("/api/pen/chat")
async def pen_chat(
    request: Request,
    show_tool_output: bool = Query(True),
    model: str = Query("gemini-3-flash-preview")
    ):
    try:
        agent = get_agent()
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))

    # Handle both JSON and FormData
    content_type = request.headers.get("content-type", "")
    message = ""
    files_data = []
    
    if "multipart/form-data" in content_type:
        form = await request.form()
        message = form.get("message", "")
        
        # Process uploaded files
        for key, value in form.items():
            if key.startswith("file_") and hasattr(value, 'read') and hasattr(value, 'filename'):
                file_content = await value.read()
                file_info = {
                    "name": value.filename,
                    "type": value.content_type or "application/octet-stream",
                    "size": len(file_content),
                    "content": file_content
                }
                files_data.append(file_info)
    else:
        body = await request.json()
        message = body.get("message", "")

    try:
        # Set the model before chat if different from current
        if model and model != agent.model_name:
            agent.set_model(model)
            logger.info(f"Model changed to: {model}")
        
        # Prepare files for agent (pass raw content, agent handles formatting)
        processed_files = []
        if files_data:
            for f in files_data:
                # Save file to user_docs/uploads
                try:
                    safe_name = "".join([c for c in f["name"] if c.isalnum() or c in "._- "])
                    saved_path = doc_manager.create_document(f"uploads/{safe_name}", f["content"].decode("utf-8", errors="ignore"))
                    logger.info(f"File saved to: {saved_path}")
                except Exception as e:
                    logger.warning(f"Could not save file {f['name']}: {e}")
                    saved_path = None

                processed_files.append({
                    "name": f["name"],
                    "type": f["type"],
                    "content": f["content"],
                    "path": saved_path # Pass the saved path to agent
                })
        
        # Run chat in a thread pool to allow SSE events to be processed
        import concurrent.futures
        loop = asyncio.get_event_loop()
        with concurrent.futures.ThreadPoolExecutor() as pool:
            reply = await loop.run_in_executor(
                pool, 
                lambda: agent.chat(message, files=processed_files if processed_files else None)
            )
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=500, detail=f"Agent error: {e}")

    # Extract tool calls from the current session (L1)
    tools: List[Dict[str, Any]] = []
    try:
        # Get the last tool interactions from L1
        tool_interactions = agent.layer_manager.l1.get_all_tool_interactions()
        
        # Return all tool interactions with complete details
        if tool_interactions:
            # Return all tool interactions from this session
            tools = tool_interactions
            
            # Optionally hide tool output if show_tool_output is False
            if not show_tool_output:
                for tool in tools:
                    tool["tool_output"] = "[Output hidden - use show_tool_output=true to view]"
    except Exception as e:
        logger.warning(f"Could not retrieve tool interactions: {e}")

    # Build file info for response (so frontend can restore on undo)
    response_files = []
    for f in files_data:
        response_files.append({
            "name": f["name"],
            "type": f["type"],
            "size": f["size"]
        })

    return {
        "reply": reply,
        "tools": tools,
        "files": response_files,  # Return file metadata for undo feature
    }


@app.get("/api/pen/models")
async def get_available_models():
    """Get list of available models with API keys configured"""
    try:
        from src.models.factory import ModelFactory
        models = ModelFactory.get_available_models_with_keys()
        return {"models": models, "default": "gemini-3-flash-preview"}
    except Exception as e:
        logger.error(f"Error getting models: {e}")
        return {"models": [], "default": "gemini-3-flash-preview"}


@app.get("/api/pen/usage")
async def get_usage_stats():
    """Get API usage statistics"""
    try:
        from src.utils.usage_logger import usage_logger
        
        summary = usage_logger.get_summary()
        model_stats = usage_logger.get_model_stats()
        daily_stats = usage_logger.get_daily_stats(days=7)
        recent = usage_logger.get_recent_entries(limit=20)
        
        return {
            "summary": summary,
            "by_model": model_stats,
            "daily": daily_stats,
            "recent_entries": recent
        }
    except Exception as e:
        logger.error(f"Error getting usage stats: {e}")
        return {"error": str(e)}

@app.get("/api/pen/archived-sessions")
async def get_archived_sessions():
    """Get list of archived sessions from L2.5 (summaries only for sidebar)"""
    try:
        agent = get_agent()
        # Get all summaries from L2.5
        summaries = agent.layer_manager.l2_5.summaries
        
        # Convert to list and sort by date (newest first)
        sessions_list = []
        for session_id, summary_data in summaries.items():
            # Create a short title from summary (first 50 chars or first sentence)
            summary_text = summary_data.get("summary", "")
            title = summary_text[:60] + "..." if len(summary_text) > 60 else summary_text
            
            sessions_list.append({
                "session_id": session_id,
                "title": title,
                "summary": summary_text,
                "keywords": summary_data.get("keywords", []),
                "created_at": summary_data.get("created_at", ""),
                "message_count": summary_data.get("message_count", 0)
            })
        
        # Sort by created_at descending (newest first)
        sessions_list.sort(key=lambda x: x.get("created_at", ""), reverse=True)
        
        return {"sessions": sessions_list}
    except Exception as e:
        logger.error(f"Error getting archived sessions: {e}", exc_info=True)
        return {"sessions": [], "error": str(e)}


@app.get("/api/pen/archived-session/{session_id}")
async def get_archived_session(session_id: str):
    """Get full archived session data from L2"""
    try:
        agent = get_agent()
        # Get full session from L2
        session = agent.layer_manager.l2.get_session_by_id(session_id)
        
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
        
        # Get summary from L2.5
        summary_data = agent.layer_manager.l2_5.get_summary_by_session_id(session_id)
        
        return {
            "session_id": session_id,
            "messages": session.get("messages", []),
            "tool_interactions": session.get("tool_interactions", []),
            "created_at": session.get("created_at", ""),
            "archived_at": session.get("archived_at", ""),
            "summary": summary_data.get("summary", "") if summary_data else "",
            "keywords": summary_data.get("keywords", []) if summary_data else []
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting archived session: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/pen/history")
async def pen_history():
    """Get chat history from L1.json"""
    try:
        logger.info("History endpoint called")
        agent = get_agent()
        session_context = agent.layer_manager.l1.get_session_context()
        messages = session_context.get("messages", [])
        tools = session_context.get("tool_interactions", [])
        logger.info(f"Returning {len(messages)} messages and {len(tools)} tool interactions")
        return {
            "messages": messages,
            "tool_interactions": tools,
        }
    except Exception as e:
        logger.error(f"Could not retrieve history: {e}", exc_info=True)
        return {
            "messages": [],
            "tool_interactions": [],
        }


@app.get("/api/pen/file-content")
async def get_file_content(name: str = Query(..., description="Relative path to the file")):
    """Get file content directly from disk (PEN WorkSpace)"""
    try:
        # Sanitize and construct path
        # Remove leading slashes to ensure it joins correctly relative to base_dir
        clean_name = name.lstrip("/\\")
        file_path = (doc_manager.base_dir / clean_name).resolve()

        # Security Check: Prevent Path Traversal
        base_path = doc_manager.base_dir.resolve()
        if not str(file_path).startswith(str(base_path)):
             logger.warning(f"Access denied: Attempted path traversal to {file_path}")
             raise HTTPException(status_code=403, detail="Access denied")

        if not file_path.exists():
            logger.warning(f"File not found: {file_path}")
            return {"content": "", "exists": False, "error": "File not found"}

        if not file_path.is_file():
             return {"content": "", "exists": False, "error": "Not a file"}

        # Determine if text or binary
        is_image = file_path.suffix.lower() in ['.png', '.jpg', '.jpeg', '.gif', '.webp', '.svg']
        
        if is_image:
            return {"content": f"[Image File: {clean_name}]", "is_image": True, "exists": True}

        try:
            content = file_path.read_text(encoding="utf-8")
            return {
                "content": content,
                "is_image": False,
                "type": file_path.suffix.lstrip("."),
                "exists": True,
                "path": clean_name
            }
        except UnicodeDecodeError:
            return {"content": "[Binary/Non-text File]", "is_image": False, "exists": True}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error reading file content: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/pen/reset")
async def pen_reset():
    try:
        agent = get_agent()
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))

    agent.reset()
    return {"status": "ok"}


@app.post("/api/pen/undo")
async def pen_undo(message_index: int = Query(..., description="Index of the user message to undo")):
    """Undo a message - removes the message and everything after it from L1"""
    try:
        agent = get_agent()
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))

    try:
        # Get the message content before removing
        l1_messages = agent.layer_manager.l1.get_all_messages()
        
        # Map frontend index (user/assistant only) to absolute L1 index
        filtered_messages = []
        mapping = {} # filtered_index -> absolute_index
        current_filtered_index = 0
        
        for i, msg in enumerate(l1_messages):
            role = msg.get("role")
            if role in ["user", "assistant"]:
                mapping[current_filtered_index] = i
                filtered_messages.append(msg)
                current_filtered_index += 1
                
        # Validate against filtered list
        if message_index < 0 or message_index >= len(filtered_messages):
            logger.warning(f"Undo index out of bounds: {message_index} (Filtered count: {len(filtered_messages)})")
            # Fallback: Try absolute index if within bounds (in case frontend logic changes)
            if 0 <= message_index < len(l1_messages):
                logger.info(f"Fallback: Using absolute index {message_index}")
                absolute_index = message_index
            else:
                raise HTTPException(status_code=400, detail="Invalid message index")
        else:
            absolute_index = mapping[message_index]
            
        target_message = l1_messages[absolute_index]
        
        if target_message.get("role") != "user":
            raise HTTPException(status_code=400, detail=f"Can only undo user messages (Found: {target_message.get('role')})")
        
        message_content = target_message.get("content", "")
        
        # Remove messages from absolute index onwards in L1
        result = agent.layer_manager.l1.undo_from_index(absolute_index)
        
        # Also remove from agent's in-memory messages
        # Find corresponding index in agent.messages (may differ due to tool messages)
        agent_msg_index = 0
        l1_user_count = 0
        
        # Count user messages up to the undo point to sync with agent.messages
        target_user_count = 0
        for i in range(absolute_index):
             if l1_messages[i].get("role") == "user":
                 target_user_count += 1
                 
        current_user_count = 0
        for i, msg in enumerate(agent.messages):
            msg_role = msg.get("role") if isinstance(msg, dict) else getattr(msg, "role", None)
            # Check for user message
            if msg_role == "user":
                if current_user_count == target_user_count:
                    agent_msg_index = i
                    break
                current_user_count += 1
        
        # Truncate agent messages
        if agent_msg_index > 0:
            agent.messages = agent.messages[:agent_msg_index]
        
        logger.info(f"Undo successful: removed messages from index {absolute_index} (Filtered: {message_index})")
        
        return {
            "status": "ok",
            "message_content": message_content,
            "removed_count": result.get("removed_count", 0)
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Undo failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Undo failed: {str(e)}")


@app.get("/api/pen/events")
async def pen_events(request: Request):
    """SSE endpoint for real-time tool activity events"""
    async def event_generator():
        queue = await event_broadcaster.subscribe()
        try:
            # Send initial connection event
            yield f"data: {json.dumps({'type': 'connected', 'timestamp': datetime.now().isoformat()})}\n\n"
            
            while True:
                # Check if client disconnected
                if await request.is_disconnected():
                    logger.info("SSE client disconnected")
                    break
                
                try:
                    # Wait for events with shorter timeout for faster disconnect detection
                    event = await asyncio.wait_for(queue.get(), timeout=15.0)
                    yield f"data: {json.dumps(event)}\n\n"
                except asyncio.TimeoutError:
                    # Send keepalive
                    yield f"data: {json.dumps({'type': 'keepalive'})}\n\n"
        except (asyncio.CancelledError, GeneratorExit):
            logger.info("SSE connection cancelled")
        finally:
            event_broadcaster.unsubscribe(queue)
            logger.info("SSE connection cleaned up")
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
            "X-Content-Type-Options": "nosniff"
        }
    )


@app.get("/api/pen/settings")
async def get_settings():
    """Get current settings including API keys (masked) and tool permissions"""
    try:
        import os
        from src.config import (
            LLM_CONFIG_OBJ, EMAIL_CONFIG_OBJ, DATA_SOURCE_CONFIG
        )
        
        # Mask API keys (show only last 4 chars)
        def mask_key(key: str) -> dict:
            if not key:
                return {"configured": False, "masked": ""}
            return {"configured": True, "masked": f"...{key[-4:]}"}
        
        return {
            "api_keys": {
                "gemini": mask_key(LLM_CONFIG_OBJ.gemini_api_key or ""),
                "openai": mask_key(LLM_CONFIG_OBJ.openai_api_key or ""),
                "anthropic": mask_key(os.getenv("ANTHROPIC_API_KEY", "")),
                "brave": mask_key(os.getenv("BRAVE_API_KEY", "")),
            },
            "tools": {
                "whatsapp": DATA_SOURCE_CONFIG.whatsapp_enabled,
                "email": EMAIL_CONFIG_OBJ.enabled,
                "drive": True,  # Always enabled if service account exists
                "web": bool(os.getenv("BRAVE_API_KEY")),
                "memory": True,  # Always enabled
            },
            "email_config": {
                "address": EMAIL_CONFIG_OBJ.address or "",
                "imap_server": EMAIL_CONFIG_OBJ.imap_server or "",
                "imap_port": EMAIL_CONFIG_OBJ.imap_port or 993,
                "configured": EMAIL_CONFIG_OBJ.is_configured(),
            }
        }
    except Exception as e:
        logger.error(f"Error getting settings: {e}")
        return {"error": str(e)}


@app.get("/api/pen/l4-profile")
async def get_l4_profile():
    """Get L4 user profile data (what the model knows about the user)"""
    try:
        agent = get_agent()
        l4 = agent.layer_manager.l4
        
        # Get full profile
        profile = l4.get_profile_for_context()
        
        # Get contexts summary
        contexts = l4.get_contexts_summary()
        
        # Get profile summary if exists
        summary = profile.get("profile_summary", "")
        
        return {
            "profile": profile,
            "contexts": contexts,
            "summary": summary,
            "metadata": l4.load_profile().get("metadata", {})
        }
    except Exception as e:
        logger.error(f"Error getting L4 profile: {e}")
        return {"error": str(e)}


@app.post("/api/pen/sync-data")
async def sync_data(source: str = Query("all")):
    """Run data sync pipeline (WhatsApp from Drive, Email if enabled)"""
    import sys
    from pathlib import Path
    from datetime import datetime
    import asyncio
    
    project_root = Path(__file__).parent
    sys.path.insert(0, str(project_root))
    
    from src.config import (
        DATA_DIR, LOGS_DIR, EMAIL_CONFIG_OBJ, SYSTEM_CONFIG,
        GOOGLE_DRIVE_CONFIG, DATA_SOURCE_CONFIG
    )
    from src.parsers.whatsapp_parser import WhatsAppParser
    from src.parsers.email_parser import EmailParser
    from src.parsers.drive_sync import auto_sync_from_drive
    from src.storage.data_manager import DataManager
    
    logs = []
    
    def add_log(message: str, level: str = "info"):
        log_entry = {"message": message, "level": level, "time": datetime.now().isoformat()}
        logs.append(log_entry)
        logger.info(f"[SYNC] {message}")
        # Broadcast sync log event
        event_broadcaster.broadcast_sync("sync_log", log_entry)
    
    def run_sync_logic():
        try:
            add_log(f"Veri guncelleme basladi (Kaynak: {source})...")
            
            data_manager = DataManager(DATA_DIR)
            
            # === WHATSAPP SYNC ===
            if (source == "all" or source == "whatsapp"):
                if DATA_SOURCE_CONFIG.whatsapp_enabled:
                    add_log("WhatsApp senkronizasyonu basliyor...")
                    
                    # Check service account
                    if not GOOGLE_DRIVE_CONFIG.service_account_file.exists():
                        add_log("Service account dosyasi bulunamadi, Drive senkronizasyonu atlaniyor", "warning")
                    else:
                        # Sync from Drive
                        whatsapp_dir = project_root / "whatsapp_export"
                        whatsapp_dir.mkdir(parents=True, exist_ok=True)
                        
                        try:
                            downloaded_files = auto_sync_from_drive(
                                service_account_file=str(GOOGLE_DRIVE_CONFIG.service_account_file),
                                output_dir=whatsapp_dir,
                                folder_name=GOOGLE_DRIVE_CONFIG.folder_name
                            )
                            
                            if downloaded_files:
                                add_log(f"Drive'dan {len(downloaded_files)} dosya indirildi", "success")
                            else:
                                add_log("Drive'da yeni dosya yok")
                        except Exception as e:
                            add_log(f"Drive senkronizasyon hatasi: {e}", "error")
                    
                    # Process WhatsApp files
                    whatsapp_dir = project_root / "whatsapp_export"
                    if whatsapp_dir.exists():
                        txt_files = list(whatsapp_dir.glob("*.txt"))
                        
                        if txt_files:
                            add_log(f"{len(txt_files)} WhatsApp dosyasi isleniyor...")
                            parser = WhatsAppParser()
                            processed_count = 0
                            
                            for txt_file in txt_files:
                                try:
                                    messages = parser.parse_file(str(txt_file))
                                    if messages:
                                        chat_name = txt_file.stem
                                        data_manager.save_whatsapp_messages(messages, chat_name)
                                        processed_count += 1
                                        # Optional: Log progress for large numbers of files
                                        if processed_count % 5 == 0:
                                            add_log(f"{processed_count}/{len(txt_files)} sohbet islendi...")
                                except Exception as e:
                                    add_log(f"Dosya isleme hatasi ({txt_file.name}): {e}", "error")
                            
                            add_log(f"{processed_count} WhatsApp sohbeti guncellendi", "success")
                        else:
                            add_log("whatsapp_export klasorunde .txt dosyasi yok", "warning")
                    else:
                        add_log("whatsapp_export klasoru bulunamadi", "warning")
                else:
                    if source == "whatsapp":
                        add_log("WhatsApp isleme devre disi (Ayarlardan etkinlestirin)", "warning")
            
            # === EMAIL SYNC ===
            if (source == "all" or source == "email"):
                if EMAIL_CONFIG_OBJ.enabled and DATA_SOURCE_CONFIG.email_enabled:
                    add_log("Email senkronizasyonu basliyor...")
                    
                    if not EMAIL_CONFIG_OBJ.is_configured():
                        add_log("Email yapilandirmasi eksik", "warning")
                    else:
                        try:
                            parser = EmailParser(
                                email_address=EMAIL_CONFIG_OBJ.address,
                                password=EMAIL_CONFIG_OBJ.password,
                                imap_server=EMAIL_CONFIG_OBJ.imap_server,
                                imap_port=EMAIL_CONFIG_OBJ.imap_port,
                                max_workers=SYSTEM_CONFIG.max_workers
                            )
                            
                            if parser.connect():
                                emails = parser.fetch_emails(folder='INBOX', limit=None, parallel=True)
                                if emails:
                                    data_manager.save_emails(emails)
                                    add_log(f"{len(emails)} email guncellendi", "success")
                                else:
                                    add_log("Yeni email yok")
                                parser.disconnect()
                            else:
                                add_log("Email sunucusuna baglanilamadi", "error")
                        except Exception as e:
                            add_log(f"Email hatasi: {e}", "error")
                else:
                    if source == "email":
                        add_log("Email isleme devre disi (Ayarlardan etkinlestirin)", "warning")
            
            # Get final statistics
            stats = data_manager.get_statistics()
            add_log(f"Durum: {stats.whatsapp_total_chats} sohbet, {stats.whatsapp_total_messages} mesaj, {stats.email_total_count} email", "success")
            add_log("Islem tamamlandi!", "success")
            
            return {
                "status": "success",
                "logs": logs,
                "statistics": {
                    "whatsapp_chats": stats.whatsapp_total_chats,
                    "whatsapp_messages": stats.whatsapp_total_messages,
                    "emails": stats.email_total_count
                }
            }
            
        except Exception as e:
            logger.error(f"Data sync failed: {e}", exc_info=True)
            add_log(f"Kritik hata: {e}", "error")
            return {
                "status": "error",
                "logs": logs,
                "error": str(e)
            }

    # Run the synchronous logic in a thread pool
    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(None, run_sync_logic)


@app.post("/api/pen/sleep")
async def pen_sleep():
    """Trigger sleep cycle to process and consolidate memories"""
    try:
        agent = get_agent()
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))

    try:
        logger.info("Sleep cycle triggered from web interface")
        # Run complete sleep cycle (includes L4 update, L1→L2+L2.5, reset)
        agent.sleep()
        logger.info("Sleep cycle completed successfully")
        return {
            "status": "success",
            "message": "Sleep cycle completed successfully. Session archived and agent reset."
        }
    except Exception as e:
        logger.error(f"Sleep cycle failed: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Sleep cycle failed: {str(e)}"
        )


class WorkspaceWriteRequest(BaseModel):
    path: str
    content: str


@app.post("/api/pen/workspace/write")
async def workspace_write(request: WorkspaceWriteRequest):
    """Write content to a file in the workspace"""
    try:
        logger.info(f"[WORKSPACE WRITE] Received request:")
        logger.info(f"  Path: {request.path}")
        logger.info(f"  Content length: {len(request.content) if request.content else 0}")
        logger.info(f"  Content preview: {request.content[:100] if request.content else 'EMPTY'}")
        
        # Validate inputs
        if not request.path:
            raise HTTPException(status_code=400, detail="Path is required")
        
        if request.content is None:
            raise HTTPException(status_code=400, detail="Content is required")
        
        # Use document manager to save file
        saved_path = doc_manager.create_document(request.path, request.content)
        
        # Verify file was actually written
        full_path = doc_manager.base_dir / saved_path
        if full_path.exists():
            actual_size = full_path.stat().st_size
            logger.info(f"[WORKSPACE WRITE] File saved successfully:")
            logger.info(f"  Saved path: {saved_path}")
            logger.info(f"  Full path: {full_path}")
            logger.info(f"  File size: {actual_size} bytes")
        else:
            logger.error(f"[WORKSPACE WRITE] File not found after save: {full_path}")
        
        return {
            "status": "success",
            "path": str(saved_path),
            "message": f"File saved: {request.path}",
            "full_path": str(full_path)
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[WORKSPACE WRITE] Error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")


@app.get("/api/pen/notes/list")
async def list_notes(category: str = Query("all")):
    """List all notes from user_docs directory"""
    try:
        # Get all documents
        all_docs = doc_manager.list_documents("")
        
        # Filter by category if specified
        if category != "all":
            all_docs = [doc for doc in all_docs if doc["path"].startswith(f"{category}/")]
        
        # Convert to notes format
        notes = []
        for doc in all_docs:
            # Skip non-markdown/text files
            if doc["type"] not in [".md", ".txt"]:
                continue
            
            # Extract title from filename
            title = Path(doc["name"]).stem.replace("-", " ").replace("_", " ").title()
            
            # Get preview (first 100 chars of content)
            try:
                content = doc_manager.read_document(doc["path"])
                preview = content[:100].replace("\n", " ").strip()
            except:
                preview = ""
            
            # Determine category from path
            path_parts = doc["path"].split("/")
            doc_category = path_parts[0] if len(path_parts) > 1 else "notes"
            
            notes.append({
                "path": doc["path"],
                "title": title,
                "preview": preview,
                "category": doc_category,
                "modified": doc["modified"],
                "size": doc["size"]
            })
        
        # Sort by modified date (newest first)
        notes.sort(key=lambda x: x["modified"], reverse=True)
        
        return {
            "status": "success",
            "notes": notes,
            "count": len(notes)
        }
    except Exception as e:
        logger.error(f"Error listing notes: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/pen/notes/{path:path}")
async def delete_note(path: str):
    """Delete a note"""
    try:
        success = doc_manager.delete_document(path)
        if success:
            return {"status": "success", "message": "Note deleted"}
        else:
            raise HTTPException(status_code=404, detail="Note not found")
    except Exception as e:
        logger.error(f"Error deleting note: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/{full_path:path}")
async def serve_static(full_path: str):
    """Serve static files and fallback to index.html for SPA routing"""
    file_path = WEB_DIR / full_path
    
    # If it's a file that exists, serve it
    if file_path.is_file():
        return FileResponse(file_path)
    
    # Otherwise serve index.html (for SPA routing)
    index_path = WEB_DIR / "index.html"
    if index_path.exists():
        return FileResponse(index_path)
    
    raise HTTPException(status_code=404, detail="Not found")


if __name__ == "__main__":
    import uvicorn
    import signal
    import sys
    
    # Graceful shutdown handler
    def signal_handler(sig, frame):
        print("\n\n🛑 Shutting down gracefully...")
        sys.exit(0)
    
    signal.signal(signal.SIGINT, signal_handler)
    signal.signal(signal.SIGTERM, signal_handler)
    
    # Run with optimized reload settings
    uvicorn.run(
        "web_server:app",
        host="127.0.0.1",
        port=8000,
        reload=True,
        reload_delay=0.5,  # Delay before reloading (prevents multiple reloads)
        timeout_graceful_shutdown=2,  # Force shutdown after 2 seconds
        log_level="info"
    )
